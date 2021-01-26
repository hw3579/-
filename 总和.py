
import docx
from docx import Document
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import StaleElementReferenceException
from bs4 import BeautifulSoup
import re
a=1
while(a>0):
 path1=input("请输入路径：")


 path=re.sub('"','',path1)

 file=docx.Document(path)
 print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
 
#输出每一段的内容
 for para in file.paragraphs:
  print(para.text)
 
#输出段落编号及段落内容
 for i in range(len(file.paragraphs)):
  print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)




       

 driver=webdriver.Chrome()
#准备url
 url='http://www.baidu.com'
 # 访问
 driver.get(url)
#浏览器大小
 driver.set_window_size(800,600)

# 找到输入框 （kw后面会具体解释，还有下面具体的标签）
 shuru=driver.find_element_by_id('kw')
# 在输入框中放我们制定的文字
 for i in range(len(file.paragraphs)):
    try:
     shuru.send_keys(file.paragraphs[i].text)
   #为了我们肉眼可见，这里休眠2秒，再进行下面的操作
     time.sleep(1)
    except Exception as e:
     print('段落是空的！')

    login=driver.find_element_by_id('su')# 定位到登录按钮的位置
    login.click()# 点击登录
    time.sleep(1)
    try:
      shuru.clear()
    except Exception as e:
     print('输入框本来就是空的')








print("没啦")



